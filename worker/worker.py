"""
DocMS Temporal Worker — Document-to-Text Extraction Pipeline

Workflows:
  - DocumentProcessingWorkflow: orchestrates extract → index

Activities:
  - extract_text: converts file bytes to plain text (docx, pdf, csv, txt, images, etc.)
  - index_to_elasticsearch: pushes extracted text into ES with classification
  - update_document_status: updates backend DB via API
"""

import asyncio
import csv
import io
import json
import logging
import mimetypes
import os
import sys
import tempfile
from dataclasses import dataclass
from datetime import timedelta
from pathlib import Path

import httpx
from elasticsearch import Elasticsearch
from temporalio import activity, workflow
from temporalio.client import Client as TemporalClient
from temporalio.worker import Worker

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(name)s — %(message)s")
logger = logging.getLogger("docms-worker")

# ── Config ────────────────────────────────────────────────────
ES_URL = os.getenv("ELASTICSEARCH_URL", "http://elasticsearch:9200")
BACKEND_URL = os.getenv("BACKEND_URL", "http://backend:8080")
TEMPORAL_HOST = os.getenv("TEMPORAL_HOST", "temporal:7233")
TASK_QUEUE = "document-processing"
ES_INDEX = "documents"
SHARED_UPLOAD_DIR = os.getenv("UPLOAD_DIR", "/data/uploads")


# ── Data Classes ──────────────────────────────────────────────
@dataclass
class DocumentTask:
    """Input for the processing workflow — passed from backend/frontend."""
    document_id: str
    user_id: str
    filename: str
    file_path: str
    content_type: str
    classification: str  # "public" or "private"


@dataclass
class ExtractionResult:
    """Output of the extract_text activity."""
    document_id: str
    extracted_text: str
    char_count: int
    success: bool
    error: str = ""


@dataclass
class IndexResult:
    """Output of index_to_elasticsearch activity."""
    document_id: str
    indexed: bool
    error: str = ""


# ── Text Extraction Activities ────────────────────────────────

@activity.defn
async def extract_text(task: DocumentTask) -> ExtractionResult:
    """
    Extracts plain text from a document file.
    Supports: .txt, .md, .csv, .tsv, .json, .xml, .html,
              .docx, .pdf, .png/.jpg/.tiff (OCR), .xlsx, .rtf
    """
    logger.info(f"Extracting text from {task.filename} ({task.content_type})")
    activity.heartbeat("starting extraction")

    file_path = Path(task.file_path)
    if not file_path.exists():
        return ExtractionResult(
            document_id=task.document_id,
            extracted_text="",
            char_count=0,
            success=False,
            error=f"File not found: {task.file_path}",
        )

    try:
        text = await _do_extract(file_path, task.content_type, task.filename)
        text = text.strip()
        logger.info(f"Extracted {len(text)} chars from {task.filename}")
        activity.heartbeat("extraction complete")
        return ExtractionResult(
            document_id=task.document_id,
            extracted_text=text,
            char_count=len(text),
            success=True,
        )
    except Exception as e:
        logger.error(f"Extraction failed for {task.filename}: {e}")
        return ExtractionResult(
            document_id=task.document_id,
            extracted_text="",
            char_count=0,
            success=False,
            error=str(e),
        )


async def _do_extract(file_path: Path, content_type: str, filename: str) -> str:
    """Route to the correct extractor based on content type."""
    ext = file_path.suffix.lower()

    # ── Plain text family ──
    if content_type in (
        "text/plain", "text/markdown", "text/rtf",
        "application/json", "application/xml",
    ) or ext in (".txt", ".md", ".rtf", ".json", ".xml", ".log", ".ini", ".cfg", ".yaml", ".yml", ".toml"):
        return _extract_plaintext(file_path)

    # ── HTML ──
    if content_type == "text/html" or ext in (".html", ".htm"):
        return _extract_html(file_path)

    # ── CSV / TSV ──
    if content_type in ("text/csv", "text/tab-separated-values") or ext in (".csv", ".tsv"):
        return _extract_csv(file_path, delimiter="\t" if ext == ".tsv" else ",")

    # ── DOCX ──
    if ext == ".docx" or "wordprocessingml" in content_type:
        return _extract_docx(file_path)

    # ── PDF ──
    if content_type == "application/pdf" or ext == ".pdf":
        return _extract_pdf(file_path)

    # ── Excel ──
    if ext in (".xlsx", ".xls") or "spreadsheetml" in content_type or "ms-excel" in content_type:
        return _extract_xlsx(file_path)

    # ── Images (OCR) ──
    if content_type.startswith("image/") or ext in (".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".gif", ".webp"):
        return _extract_image_ocr(file_path)

    # ── Fallback: try as text ──
    logger.warning(f"Unknown type {content_type} ({ext}), trying plaintext fallback")
    return _extract_plaintext(file_path)


# ── Extractor Implementations ─────────────────────────────────

def _extract_plaintext(path: Path) -> str:
    """Read raw text with encoding detection fallback."""
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return path.read_text(encoding=enc)
        except (UnicodeDecodeError, ValueError):
            continue
    return path.read_bytes().decode("utf-8", errors="replace")


def _extract_html(path: Path) -> str:
    """Strip HTML tags to get text content."""
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="replace"), "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)
    except ImportError:
        import re
        raw = path.read_text(encoding="utf-8", errors="replace")
        return re.sub(r"<[^>]+>", " ", raw)


def _extract_csv(path: Path, delimiter: str = ",") -> str:
    """Convert CSV/TSV rows into searchable text."""
    lines = []
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f, delimiter=delimiter)
        for i, row in enumerate(reader):
            if i == 0:
                lines.append("Columns: " + " | ".join(row))
            else:
                lines.append(" | ".join(row))
            if i > 500:
                lines.append(f"... ({i}+ rows total)")
                break
    return "\n".join(lines)


def _extract_docx(path: Path) -> str:
    """Extract text from .docx files."""
    try:
        import docx
        doc = docx.Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n\n".join(paragraphs)
    except ImportError:
        # Fallback: docx is a zip with XML inside
        import zipfile
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile(path) as z:
            with z.open("word/document.xml") as f:
                tree = ET.parse(f)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        texts = [node.text for node in tree.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t") if node.text]
        return " ".join(texts)


def _extract_pdf(path: Path) -> str:
    """Extract text from PDF files."""
    try:
        import pymupdf  # PyMuPDF / fitz
        doc = pymupdf.open(str(path))
        pages = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                pages.append(f"--- Page {page_num + 1} ---\n{text}")
            else:
                # Page has no text — try OCR on the rendered image
                pix = page.get_pixmap(dpi=200)
                img_bytes = pix.tobytes("png")
                ocr_text = _ocr_bytes(img_bytes)
                if ocr_text.strip():
                    pages.append(f"--- Page {page_num + 1} (OCR) ---\n{ocr_text}")
        doc.close()
        return "\n\n".join(pages) if pages else "(empty PDF)"
    except ImportError:
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"--- Page {i+1} ---\n{text}")
            return "\n\n".join(pages) if pages else "(could not extract text from PDF)"
        except ImportError:
            return "(PDF extraction libraries not available)"


def _extract_xlsx(path: Path) -> str:
    """Extract text from Excel files."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        sheets = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    rows.append(" | ".join(cells))
                if i > 500:
                    rows.append(f"... ({i}+ rows)")
                    break
            if rows:
                sheets.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows))
        wb.close()
        return "\n\n".join(sheets)
    except ImportError:
        return "(Excel extraction library not available)"


def _extract_image_ocr(path: Path) -> str:
    """OCR an image file using Tesseract."""
    return _ocr_bytes(path.read_bytes())


def _ocr_bytes(img_bytes: bytes) -> str:
    """Run OCR on raw image bytes."""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(img_bytes))
        text = pytesseract.image_to_string(img)
        return text.strip() if text.strip() else "(no text detected in image)"
    except ImportError:
        return "(OCR libraries not available — install pytesseract and Pillow)"
    except Exception as e:
        return f"(OCR error: {e})"


# ── Elasticsearch Indexing Activity ───────────────────────────

@activity.defn
async def index_to_elasticsearch(task: DocumentTask, extraction: ExtractionResult) -> IndexResult:
    """Index the extracted text into Elasticsearch with classification."""
    logger.info(f"Indexing document {task.document_id} into ES (classification={task.classification})")
    activity.heartbeat("indexing to elasticsearch")

    if not extraction.success or not extraction.extracted_text:
        return IndexResult(
            document_id=task.document_id,
            indexed=False,
            error="No text to index",
        )

    try:
        es = Elasticsearch(ES_URL)

        doc_body = {
            "doc_id": task.document_id,
            "user_id": task.user_id,
            "filename": task.filename,
            "classification": task.classification,
            "content": extraction.extracted_text,
            "content_type": task.content_type,
            "indexed_at": "now",
        }

        es.index(index=ES_INDEX, id=task.document_id, document=doc_body, refresh="true")
        logger.info(f"Document {task.document_id} indexed successfully")

        return IndexResult(document_id=task.document_id, indexed=True)
    except Exception as e:
        logger.error(f"ES indexing failed: {e}")
        return IndexResult(document_id=task.document_id, indexed=False, error=str(e))


# ── Status Update Activity ────────────────────────────────────

@activity.defn
async def update_document_status(document_id: str, status: str, extracted_text: str) -> bool:
    """Notify the backend of the processing result."""
    logger.info(f"Updating document {document_id} status to '{status}'")
    activity.heartbeat("updating status")
    try:
        async with httpx.AsyncClient() as client:
            resp = await client.post(
                f"{BACKEND_URL}/internal/update-status",
                json={
                    "document_id": document_id,
                    "status": status,
                    "extracted_text": extracted_text[:50000],  # cap at 50k
                },
                timeout=10,
            )
            return resp.status_code == 200
    except Exception as e:
        logger.warning(f"Status update failed (non-fatal): {e}")
        return False


# ── Workflow Definition ───────────────────────────────────────

@workflow.defn
class DocumentProcessingWorkflow:
    """
    Temporal workflow: Extract text → Index to ES → Update status.

    Called by the Go backend when a file is uploaded, or directly
    from the frontend via the /api/temporal/start endpoint.
    """

    @workflow.run
    async def run(self, task: DocumentTask) -> dict:
        workflow.logger.info(f"Starting workflow for {task.filename} (doc={task.document_id})")

        # Step 1 — Extract text
        extraction: ExtractionResult = await workflow.execute_activity(
            extract_text,
            task,
            start_to_close_timeout=timedelta(minutes=10),
            heartbeat_timeout=timedelta(seconds=60),
        )

        if not extraction.success:
            await workflow.execute_activity(
                update_document_status,
                args=[task.document_id, "failed", extraction.error],
                start_to_close_timeout=timedelta(seconds=30),
            )
            return {
                "document_id": task.document_id,
                "status": "failed",
                "error": extraction.error,
            }

        # Step 2 — Index into Elasticsearch
        index_result: IndexResult = await workflow.execute_activity(
            index_to_elasticsearch,
            args=[task, extraction],
            start_to_close_timeout=timedelta(minutes=2),
            heartbeat_timeout=timedelta(seconds=30),
        )

        final_status = "completed" if index_result.indexed else "index_failed"

        # Step 3 — Update backend status
        await workflow.execute_activity(
            update_document_status,
            args=[task.document_id, final_status, extraction.extracted_text],
            start_to_close_timeout=timedelta(seconds=30),
        )

        return {
            "document_id": task.document_id,
            "status": final_status,
            "chars_extracted": extraction.char_count,
            "indexed": index_result.indexed,
        }


# ── Worker Entry Point ────────────────────────────────────────

async def run_worker():
    logger.info("═══════════════════════════════════════════")
    logger.info("  DocMS Temporal Worker")
    logger.info(f"  Temporal: {TEMPORAL_HOST}")
    logger.info(f"  ES:       {ES_URL}")
    logger.info(f"  Backend:  {BACKEND_URL}")
    logger.info(f"  Queue:    {TASK_QUEUE}")
    logger.info("═══════════════════════════════════════════")

    # Wait for Temporal to be ready
    client = None
    for attempt in range(60):
        try:
            client = await TemporalClient.connect(TEMPORAL_HOST)
            logger.info("Connected to Temporal server")
            break
        except Exception as e:
            if attempt % 10 == 0:
                logger.info(f"Waiting for Temporal server... (attempt {attempt + 1})")
            await asyncio.sleep(2)

    if not client:
        logger.error("Could not connect to Temporal after 120s")
        sys.exit(1)

    # Wait for Elasticsearch
    for attempt in range(30):
        try:
            es = Elasticsearch(ES_URL)
            es.info()
            logger.info("Elasticsearch is reachable")
            break
        except Exception:
            if attempt % 5 == 0:
                logger.info(f"Waiting for Elasticsearch... (attempt {attempt + 1})")
            await asyncio.sleep(2)

    worker = Worker(
        client,
        task_queue=TASK_QUEUE,
        workflows=[DocumentProcessingWorkflow],
        activities=[extract_text, index_to_elasticsearch, update_document_status],
    )

    logger.info(f"Worker listening on queue '{TASK_QUEUE}'")
    await worker.run()


if __name__ == "__main__":
    asyncio.run(run_worker())
