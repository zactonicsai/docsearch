"""
Activities for document processing.

These run outside the Temporal workflow sandbox, so heavy imports
(httpx, elasticsearch, PIL, pymupdf, etc.) are safe here.
"""

import csv
import io
import logging
import os
from pathlib import Path

import httpx
from elasticsearch import Elasticsearch
from temporalio import activity

from shared import DocumentTask, ExtractionResult, IndexResult

logger = logging.getLogger("docms-activities")

ES_URL = os.getenv("ELASTICSEARCH_URL", "http://elasticsearch:9200")
BACKEND_URL = os.getenv("BACKEND_URL", "http://backend:8080")
ES_INDEX = "documents"


# ── Extract Text Activity ────────────────────────────────────

@activity.defn
async def extract_text(task: DocumentTask) -> ExtractionResult:
    """Extract plain text from a document file."""
    logger.info(f"Extracting text from {task.filename} ({task.content_type})")
    activity.heartbeat("starting extraction")

    file_path = Path(task.file_path)
    if not file_path.exists():
        return ExtractionResult(
            document_id=task.document_id, extracted_text="",
            char_count=0, success=False,
            error=f"File not found: {task.file_path}",
        )

    try:
        text = _do_extract(file_path, task.content_type, task.filename)
        text = text.strip()
        logger.info(f"Extracted {len(text)} chars from {task.filename}")
        activity.heartbeat("extraction complete")
        return ExtractionResult(
            document_id=task.document_id, extracted_text=text,
            char_count=len(text), success=True,
        )
    except Exception as e:
        logger.error(f"Extraction failed for {task.filename}: {e}")
        return ExtractionResult(
            document_id=task.document_id, extracted_text="",
            char_count=0, success=False, error=str(e),
        )


def _do_extract(file_path: Path, content_type: str, filename: str) -> str:
    """Route to the correct extractor."""
    ext = file_path.suffix.lower()

    # Plain text family
    if content_type in (
        "text/plain", "text/markdown", "text/rtf",
        "application/json", "application/xml",
    ) or ext in (".txt", ".md", ".rtf", ".json", ".xml", ".log", ".yaml", ".yml", ".toml"):
        return _extract_plaintext(file_path)

    if content_type == "text/html" or ext in (".html", ".htm"):
        return _extract_html(file_path)

    if content_type in ("text/csv", "text/tab-separated-values") or ext in (".csv", ".tsv"):
        return _extract_csv(file_path, delimiter="\t" if ext == ".tsv" else ",")

    if ext == ".docx" or "wordprocessingml" in content_type:
        return _extract_docx(file_path)

    if content_type == "application/pdf" or ext == ".pdf":
        return _extract_pdf(file_path)

    if ext in (".xlsx", ".xls") or "spreadsheetml" in content_type or "ms-excel" in content_type:
        return _extract_xlsx(file_path)

    if content_type.startswith("image/") or ext in (
        ".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".gif", ".webp",
    ):
        return _extract_image_ocr(file_path)

    logger.warning(f"Unknown type {content_type} ({ext}), trying plaintext fallback")
    return _extract_plaintext(file_path)


# ── Extractors ────────────────────────────────────────────────

def _extract_plaintext(path: Path) -> str:
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return path.read_text(encoding=enc)
        except (UnicodeDecodeError, ValueError):
            continue
    return path.read_bytes().decode("utf-8", errors="replace")


def _extract_html(path: Path) -> str:
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="replace"), "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        return soup.get_text(separator="\n", strip=True)
    except ImportError:
        import re
        raw = path.read_text(encoding="utf-8", errors="replace")
        return re.sub(r"<[^>]+>", " ", raw)


def _extract_csv(path: Path, delimiter: str = ",") -> str:
    lines = []
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f, delimiter=delimiter)
        for i, row in enumerate(reader):
            if i == 0:
                lines.append("Columns: " + " | ".join(row))
            else:
                lines.append(" | ".join(row))
            if i > 500:
                lines.append(f"... ({i}+ rows total)")
                break
    return "\n".join(lines)


def _extract_docx(path: Path) -> str:
    try:
        import docx
        doc = docx.Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n\n".join(paragraphs)
    except ImportError:
        import zipfile
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile(path) as z:
            with z.open("word/document.xml") as f:
                tree = ET.parse(f)
        texts = [
            node.text
            for node in tree.iter(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"
            )
            if node.text
        ]
        return " ".join(texts)


def _extract_pdf(path: Path) -> str:
    try:
        import pymupdf
        doc = pymupdf.open(str(path))
        pages = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                pages.append(f"--- Page {page_num + 1} ---\n{text}")
            else:
                pix = page.get_pixmap(dpi=200)
                ocr_text = _ocr_bytes(pix.tobytes("png"))
                if ocr_text.strip():
                    pages.append(f"--- Page {page_num + 1} (OCR) ---\n{ocr_text}")
        doc.close()
        return "\n\n".join(pages) if pages else "(empty PDF)"
    except ImportError:
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"--- Page {i + 1} ---\n{text}")
            return "\n\n".join(pages) if pages else "(could not extract text)"
        except ImportError:
            return "(PDF extraction libraries not available)"


def _extract_xlsx(path: Path) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        sheets = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                cells = [str(c) if c is not None else "" for c in row]
                if any(cells):
                    rows.append(" | ".join(cells))
                if i > 500:
                    rows.append(f"... ({i}+ rows)")
                    break
            if rows:
                sheets.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows))
        wb.close()
        return "\n\n".join(sheets)
    except ImportError:
        return "(Excel extraction library not available)"


def _extract_image_ocr(path: Path) -> str:
    return _ocr_bytes(path.read_bytes())


def _ocr_bytes(img_bytes: bytes) -> str:
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(img_bytes))
        text = pytesseract.image_to_string(img)
        return text.strip() if text.strip() else "(no text detected in image)"
    except ImportError:
        return "(OCR libraries not available)"
    except Exception as e:
        return f"(OCR error: {e})"


# ── Index to Elasticsearch Activity ──────────────────────────

@activity.defn
async def index_to_elasticsearch(task: DocumentTask, extraction: ExtractionResult) -> IndexResult:
    """Index extracted text into Elasticsearch with classification."""
    logger.info(f"Indexing {task.document_id} into ES (class={task.classification})")
    activity.heartbeat("indexing")

    if not extraction.success or not extraction.extracted_text:
        return IndexResult(document_id=task.document_id, indexed=False, error="No text")

    try:
        es = Elasticsearch(ES_URL)
        es.index(
            index=ES_INDEX,
            id=task.document_id,
            document={
                "doc_id": task.document_id,
                "user_id": task.user_id,
                "filename": task.filename,
                "classification": task.classification,
                "content": extraction.extracted_text,
                "content_type": task.content_type,
                "indexed_at": "now",
            },
            refresh="true",
        )
        logger.info(f"Document {task.document_id} indexed")
        return IndexResult(document_id=task.document_id, indexed=True)
    except Exception as e:
        logger.error(f"ES indexing failed: {e}")
        return IndexResult(document_id=task.document_id, indexed=False, error=str(e))


# ── Update Status Activity ───────────────────────────────────

@activity.defn
async def update_document_status(document_id: str, status: str, extracted_text: str) -> bool:
    """Callback to backend with processing result."""
    logger.info(f"Updating {document_id} → '{status}'")
    activity.heartbeat("updating status")
    try:
        async with httpx.AsyncClient() as client:
            resp = await client.post(
                f"{BACKEND_URL}/internal/update-status",
                json={
                    "document_id": document_id,
                    "status": status,
                    "extracted_text": extracted_text[:50000],
                },
                timeout=10,
            )
            return resp.status_code == 200
    except Exception as e:
        logger.warning(f"Status update failed (non-fatal): {e}")
        return False
